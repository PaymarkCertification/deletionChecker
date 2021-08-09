from docx import Document
import io
import glob
import pandas as pd
import numpy as np
import unittest
import datetime
name = 'DELETE - S92031MODPMK2102.docx'


class extractDocx(object):
    TBL_RQ  = 0 # Request By, Load Before, Project ID
    TBL_SG  = 1 # Snr sig, Co-ord sig, ADD, UPDATE, DELETE
    TBL_SV  = 2 # Software Version, Char count, status
    TBL_DBN = 3 # DB Notes
    TBL_CER = 4 # Certification Use, Reason for deletion, No term active
    TBL_OPS = 5 # OPs use, et al

    def __init__(self, document):
        self.doc = Document(document)
        self.software_version = self.software(self).get_software_version
        self.status           = self.software(self).get_status
        self.request_by       = self.request(self).get_request_by
        self.load_before      = self.request(self).get_load_before
        self.project_id       = self.request(self).get_project_id 
        self.reason           = self.certification(self).get_reason
        self.confirm          = self.certification(self).get_confirm

    @property
    def set_docx(self):
        return self.doc

    def __set_tables(self) -> tuple[pd.DataFrame]:
        self.tables = []
        for tab in self.doc.tables:
            df = [['' for i in range(len(tab.columns))] for j in range(len(tab.rows))]
            for i, row in enumerate(tab.rows):
                for j, cell in enumerate(row.cells):
                    if cell.text:
                        df[i][j] = cell.text
            self.tables.append(pd.DataFrame(df))
        return self.tables
    
    def __get_tables(self, t_num):
        self.__set_tables()
        return self.tables[t_num]
    
    @property
    def get_software_tbl(self) -> pd.DataFrame:
        return self.__get_tables(self.TBL_SV)

    @property
    def get_rq_tbl(self) -> pd.DataFrame:
        return self.__get_tables(self.TBL_RQ)

    @property
    def get_sg_tbl(self) -> pd.DataFrame:
        return self.__get_tables(self.TBL_SG)

    @property
    def get_dbn_tbl(self) -> pd.DataFrame:
        return self.__get_tables(self.TBL_DBN)

    @property
    def get_cer_tbl(self) -> pd.DataFrame:
        return self.__get_tables(self.TBL_CER)

    @property
    def get_ops_tbl(self) -> pd.DataFrame:
        return self.__get_tables(self.TBL_OPS)

    class software(object):
        def __init__(self, outer):
            self.outer = outer
            self.sw_tbl = self.outer.get_software_tbl

        @property
        def get_software_version(self) -> str:
            return self.outer.join_string(self.sw_tbl, 0, 2, 18)

        @property
        def get_status(self) -> str:
            return self.outer.join_string(self.sw_tbl, 0, 21, 23)

    class request(object):
        def __init__(self, outer):
            self.outer = outer
            self.rq_tbl = outer.get_rq_tbl

        @property
        def get_request_by(self) -> str:
            return self.outer.join_string(self.rq_tbl, 0, 1, 2)

        @property
        def get_load_before(self) -> str:
            return self.outer.join_string(self.rq_tbl, 0, 3, 4)

        @property
        def get_project_id(self) -> str:
            return self.outer.join_string(self.rq_tbl, 0, 5 ,6)

    class certification(object):
        def __init__(self, outer):
            self.outer = outer
            self.cer_tbl = outer.get_cer_tbl

        @property
        def get_reason(self) -> str:
            return self.outer.join_string(self.cer_tbl,1, 1, 2)

        @property
        def get_confirm(self) -> str:
            return self.outer.join_string(self.cer_tbl, 2, 2, 3)

    @staticmethod
    def get_today() -> datetime:
        return datetime.time().strftime("%d-%m-%Y")

    @staticmethod
    def join_string(table: pd.DataFrame, index: int, slice1: int, slice2: int) -> str:
        return f"{''.join(table.iloc[index].values[slice1:slice2])}"


class test_checker(unittest.TestCase):
    def setUp(self):
        self.x = extractDocx(name)

    def test_software_version_equals_filename_sv(self):
        self.assertEqual(self.x.software_version, name[9:-5])

    def test_neg_software_version_check(self):
        self.assertNotEqual(self.x.software_version, 'Incorrect Software') 

    def test_Request_by_not_empty(self):
        self.assertNotEqual(self.x.request_by, None)

    def test_load_by_date_valid(self):
        self.assertTrue(self.x.load_before>=self.x.get_today())

    def test_project_id_not_empty(self):
        self.assertNotEqual(self.x.project_id, None)

    def test_reason_for_deletion(self):
        self.assertEqual(self.x.reason, "SVT has passed Sunset and has no active users ")

    def test_confirmation_equals_yes(self):
        self.assertEqual(self.x.confirm, "Yes")

    def test_project_status_equals_00(self):
        self.assertEqual(self.x.status, "00")

    def tearDown(self):
        pass

def r_var(txt, attr):
    return f"{'-'*5}> Getting {txt}...\n \t{txt}: {attr}"


if __name__=='__main__':

    x = extractDocx(name)
    print("Loading docx Attributes...\n")
    print(r_var('Software Version', x.software_version))
    print(r_var("Request By", x.request_by))
    print(r_var("Load Before", x.load_before))
    print(r_var("Project ID", x.project_id))
    print(r_var("Status", x.status))
    print(r_var("Reason", x.reason))
    print(r_var("Confirmation", x.confirm))
    print("\nChecking Document...")
    unittest.main(verbosity=3)
